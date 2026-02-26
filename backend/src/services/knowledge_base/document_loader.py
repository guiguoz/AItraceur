# =============================================
# Chargeur de documents
# Sprint 6: Base de connaissances RAG
# =============================================

import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Union

# Import conditionnel pour PyPDF2
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Import conditionnel pour python-docx
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# =============================================
# Types de données
# =============================================
@dataclass
class Document:
    """Un document importé."""

    doc_id: str
    title: str
    source: str  # Nom du fichier ou URL
    doc_type: str  # pdf, docx, txt, url
    content: str
    created_at: datetime = field(default_factory=datetime.utcnow)
    metadata: Dict = field(default_factory=dict)


@dataclass
class DocumentChunk:
    """Un chunk de document pour le RAG."""

    chunk_id: str
    doc_id: str
    content: str
    source: str
    page: Optional[int] = None
    section: Optional[str] = None
    metadata: Dict = field(default_factory=dict)


# =============================================
# Chargeur de documents
# =============================================
class DocumentLoader:
    """
    Chargeur de documents pour la base de connaissances.

    Supporte:
    - PDF (.pdf)
    - Word (.docx)
    - Texte (.txt)
    - URLs web
    """

    CHUNK_SIZE = 1000  # Caractères par chunk
    CHUNK_OVERLAP = 200  # Chevauchement entre chunks

    def __init__(self, chunk_size: int = None):
        """
        Initialise le chargeur.

        Args:
            chunk_size: Taille des chunks en caractères
        """
        self.chunk_size = chunk_size or self.CHUNK_SIZE
        self.documents: List[Document] = []

    def load_file(self, file_path: Union[str, Path]) -> Optional[Document]:
        """
        Charge un fichier.

        Args:
            file_path: Chemin du fichier

        Returns:
            Document ou None si erreur
        """
        path = Path(file_path)

        if not path.exists():
            print(f"Fichier non trouvé: {path}")
            return None

        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._load_pdf(path)
        elif ext == ".docx":
            return self._load_docx(path)
        elif ext == ".txt":
            return self._load_txt(path)
        else:
            print(f"Type de fichier non supporté: {ext}")
            return None

    def load_url(self, url: str) -> Optional[Document]:
        """
        Charge une page web.

        Args:
            url: URL de la page

        Returns:
            Document ou None si erreur
        """
        try:
            import requests
            from bs4 import BeautifulSoup

            response = requests.get(url, timeout=10)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, "html.parser")

            # Extraire le titre
            title = soup.title.string if soup.title else url

            # Extraire le texte (enlever scripts et styles)
            for script in soup(["script", "style"]):
                script.decompose()

            content = soup.get_text(separator="\n", strip=True)

            # Créer le document
            doc = Document(
                doc_id=self._generate_id(url),
                title=title,
                source=url,
                doc_type="url",
                content=content,
                metadata={"url": url},
            )

            self.documents.append(doc)
            return doc

        except Exception as e:
            print(f"Erreur chargement URL: {e}")
            return None

    def load_directory(
        self, directory: Union[str, Path], extensions: List[str] = None
    ) -> List[Document]:
        """
        Charge tous les fichiers d'un répertoire.

        Args:
            directory: Répertoire
            extensions: Extensions à charger (ex: [".pdf", ".docx"])

        Returns:
            Liste de documents
        """
        directory = Path(directory)
        extensions = extensions or [".pdf", ".docx", ".txt"]

        documents = []

        for ext in extensions:
            for file_path in directory.rglob(f"*{ext}"):
                doc = self.load_file(file_path)
                if doc:
                    documents.append(doc)

        return documents

    def chunk_document(self, document: Document) -> List[DocumentChunk]:
        """
        Découpe un document en chunks.

        Args:
            document: Document à chunker

        Returns:
            Liste de chunks
        """
        chunks = []
        content = document.content

        # Nettoyer le contenu
        content = self._clean_text(content)

        # Découper en paragraphs
        paragraphs = content.split("\n\n")

        current_chunk = ""
        chunk_index = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Si le paragraphe est trop grand, le découpe
            if len(para) > self.chunk_size:
                if current_chunk:
                    chunks.append(
                        self._create_chunk(document, current_chunk, chunk_index)
                    )
                    chunk_index += 1
                    current_chunk = ""

                # Découper le paragraphe long
                for i in range(0, len(para), self.chunk_size - self.CHUNK_OVERLAP):
                    sub_para = para[i : i + self.chunk_size]
                    chunks.append(self._create_chunk(document, sub_para, chunk_index))
                    chunk_index += 1

            elif len(current_chunk) + len(para) > self.chunk_size:
                # Ajouter le chunk actuel et commencer un nouveau
                chunks.append(self._create_chunk(document, current_chunk, chunk_index))
                chunk_index += 1
                current_chunk = para
            else:
                # Ajouter au chunk actuel
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para

        # Ajouter le dernier chunk
        if current_chunk:
            chunks.append(self._create_chunk(document, current_chunk, chunk_index))

        return chunks

    def _load_pdf(self, path: Path) -> Optional[Document]:
        """Charge un PDF."""
        if not PDF_AVAILABLE:
            print("PyPDF2 non installé: pip install PyPDF2")
            return None

        try:
            content = ""
            page_count = 0

            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                page_count = len(reader.pages)

                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        content += f"\n--- Page {i + 1} ---\n{text}"

            doc = Document(
                doc_id=self._generate_id(str(path)),
                title=path.stem,
                source=str(path),
                doc_type="pdf",
                content=content,
                metadata={"pages": page_count},
            )

            self.documents.append(doc)
            return doc

        except Exception as e:
            print(f"Erreur lecture PDF: {e}")
            return None

    def _load_docx(self, path: Path) -> Optional[Document]:
        """Charge un fichier Word."""
        if not DOCX_AVAILABLE:
            print("python-docx non installé: pip install python-docx")
            return None

        try:
            doc = docx.Document(path)
            content = "\n\n".join([para.text for para in doc.paragraphs])

            document = Document(
                doc_id=self._generate_id(str(path)),
                title=path.stem,
                source=str(path),
                doc_type="docx",
                content=content,
                metadata={"paragraphs": len(doc.paragraphs)},
            )

            self.documents.append(document)
            return document

        except Exception as e:
            print(f"Erreur lecture DOCX: {e}")
            return None

    def _load_txt(self, path: Path) -> Optional[Document]:
        """Charge un fichier texte."""
        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()

            doc = Document(
                doc_id=self._generate_id(str(path)),
                title=path.stem,
                source=str(path),
                doc_type="txt",
                content=content,
            )

            self.documents.append(doc)
            return doc

        except Exception as e:
            print(f"Erreur lecture TXT: {e}")
            return None

    def _create_chunk(
        self, document: Document, content: str, index: int
    ) -> DocumentChunk:
        """Crée un chunk."""
        return DocumentChunk(
            chunk_id=f"{document.doc_id}_chunk_{index}",
            doc_id=document.doc_id,
            content=content,
            source=document.source,
            metadata=document.metadata,
        )

    def _clean_text(self, text: str) -> str:
        """Nettoie le texte."""
        # Supprimer les espaces multiples
        text = re.sub(r" +", " ", text)

        # Supprimer les lignes vides excessives
        text = re.sub(r"\n\n\n+", "\n\n", text)

        return text.strip()

    def _generate_id(self, source: str) -> str:
        """Génère un ID unique."""
        import hashlib

        return hashlib.md5(source.encode()).hexdigest()[:16]


# =============================================
# Documents officiels CO (modèles)
# =============================================
class OfficialDocuments:
    """
    Référentiel des documents officiels IOF et FFCO.

    Ces documents peuvent être téléchargés et indexés.
    """

    # URLs des documents officiels (exemples)
    IOF_DOCUMENTS = {
        "isom_2017": {
            "title": "International Specification for Orienteering Maps 2017",
            "url": "https://orienteering.org/wp-content/uploads/2021/02/ISOM_2017-2.pdf",
            "description": "Norme internationale pour les cartes de forêt",
        },
        "issprom_2019": {
            "title": "International Specification for Sprint Orienteering Maps 2019",
            "url": "https://orienteering.org/wp-content/uploads/2021/02/ISSprOM_2019.pdf",
            "description": "Norme internationale pour les cartes de sprint",
        },
        "ireco_2022": {
            "title": "International Rules for Elite Competitions 2022",
            "url": "https://orienteering.org/wp-content/uploads/2021/02/IRECO_2022.pdf",
            "description": "Règles officielles pour les compétitions d'élite",
        },
    }

    FFCO_DOCUMENTS = {
        "reglement_technique": {
            "title": "Règlement Technique FFCO",
            "url": "https://www.ffcorientation.fr/wp-content/uploads/reglement_technique.pdf",
            "description": "Règlement technique de la Fédération Française de CO",
        },
    }

    @classmethod
    def get_all_documents(cls) -> Dict:
        """Retourne tous les documents disponibles."""
        return {
            "iof": cls.IOF_DOCUMENTS,
            "ffco": cls.FFCO_DOCUMENTS,
        }
